"""
Generate the cooperation agreement draft between Xiamen University and Blue Ocean Group,
based on the 706 Institute reference agreement structure, focused on scientific research
cooperation on Mintou-1 (闽投一号).
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn


def set_chinese_font(run, font_name="宋体", size=12, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    r = run._element
    rPr = r.find(qn("w:rPr"))
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def add_title(doc, text, size=20):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_chinese_font(run, "黑体", size, bold=True)
    return p


def add_heading_cn(doc, text, level=1):
    sizes = {1: 16, 2: 14, 3: 13}
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, "黑体", sizes.get(level, 12), bold=True)
    return p


def add_body(doc, text, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_chinese_font(run, "宋体", 12)
    return p


def add_blank(doc):
    doc.add_paragraph()


doc = Document()

# Set default style
style = doc.styles["Normal"]
style.font.name = "宋体"
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

# Cover / title
for _ in range(3):
    add_blank(doc)

add_title(doc, "蔚蓝海洋集团有限公司", size=22)
add_title(doc, "厦门大学", size=22)
add_blank(doc)
add_title(doc, "闽投一号海洋科研合作", size=20)
add_title(doc, "框架协议", size=20)

for _ in range(8):
    add_blank(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("二〇二六年    月")
set_chinese_font(run, "宋体", 14)

doc.add_page_break()

# Parties
add_body(doc, "甲方：蔚蓝海洋集团有限公司", indent=False)
add_body(doc, "地址：", indent=False)
add_body(doc, "法定代表人：", indent=False)
add_blank(doc)
add_body(doc, "乙方：厦门大学", indent=False)
add_body(doc, "地址：福建省厦门市思明区思明南路422号", indent=False)
add_body(doc, "法定代表人：张宗益", indent=False)
add_blank(doc)

# Preamble
add_body(doc,
    "为深入贯彻落实党中央、国务院关于建设海洋强国、发展海洋经济的战略部署，"
    "积极响应国家《\"十四五\"海洋经济发展规划》《关于建立健全海洋生态产品价值实现机制的若干意见》"
    "以及福建省关于深远海养殖、海洋牧场和蓝色碳汇等相关政策导向，"
    "双方本着优势互补、合作共赢的原则，"
    "充分发挥甲方在深远海养殖装备、海洋牧场运营和海洋产业资源整合方面的产业优势，"
    "以及乙方在海洋科学、生态学、信息工程及国际科研合作方面的学科与人才优势，"
    "经友好协商，决定依托甲方战略性资产——\"闽投一号\"深远海智能养殖平台，"
    "共同推动ONCE国际大科学计划（Ocean Negative Carbon Emissions，海洋负排放国际大科学计划）"
    "在中国的落地，将闽投一号挂牌建设为ONCE国际大科学计划在华创新科研基地，"
    "围绕海洋环境智能监测、水域净化与生态修复、水下巡检与数据采集等方向开展科研合作，"
    "推动\"国际科研方法学 × 中国海洋场景\"对接验证，"
    "服务福建省海洋经济高质量发展，特签订本框架协议。")

add_blank(doc)

# Section 1: Cooperation content
add_heading_cn(doc, "一、合作内容", level=1)

add_heading_cn(doc, "（一）合作原则与共识", level=2)
add_body(doc,
    "1. 双方深刻认识加快建设海洋强国、发展海洋经济的重要意义，积极落实国家和福建省关于海洋科技创新、"
    "深远海养殖与海洋生态保护的决策部署，全方位推进闽投一号上的科研合作与成果转化。")
add_body(doc,
    "2. 充分发挥甲方依托闽投一号等核心资产形成的深远海养殖与海洋牧场产业优势，"
    "以及乙方作为国家\"双一流\"建设高校的科研平台优势（含近海海洋环境科学国家重点实验室、"
    "海洋与地球学院、生命科学学院、环境与生态学院、信息学院等），"
    "积极探索产学研用深度融合的新模式，形成优势互补、利益共享的合作共同体。")
add_body(doc,
    "3. 双方一致同意，以ONCE国际大科学计划在闽投一号挂牌为标志性事件，"
    "将闽投一号建设为兼具国际科研体系背书与中国海洋场景验证能力的合作示范平台，"
    "打造中国深远海科研合作的物理载体与品牌窗口。")

add_heading_cn(doc, "（二）ONCE国际大科学计划挂牌", level=2)
add_body(doc,
    "1. 双方共同推动将ONCE国际大科学计划"
    "（Ocean Negative Carbon Emissions，海洋负排放国际大科学计划）"
    "在中国的创新科研基地挂牌于闽投一号深远海智能养殖平台。"
    "挂牌名称拟定为\"ONCE国际大科学计划·闽投一号创新科研基地\""
    "（最终以双方书面确认为准）。")
add_body(doc,
    "2. 由甲方负责提供闽投一号平台、海域作业场地、配套设施及现场作业保障，"
    "并就挂牌仪式、品牌使用、对外宣传等事项与乙方协同推进；"
    "由乙方负责对接ONCE国际大科学计划国际秘书处与相关国际科研体系，"
    "依据ONCE计划的总体方法学与治理框架，推动其在闽投一号场景下的本地化落地与方法学验证。")
add_body(doc,
    "3. 双方就挂牌后的对外宣传、媒体披露、品牌使用、标识展示、共同对外署名等事项"
    "建立联合审定机制。任一方对外使用ONCE国际大科学计划名称、标识及挂牌相关信息，"
    "应事先取得另一方书面同意，并遵守ONCE国际大科学计划的总体品牌规范。")
add_body(doc,
    "4. 围绕挂牌基地，双方共同策划与举办挂牌仪式、揭牌活动、国际学术研讨、"
    "现场考察与海外访问团接待等品牌活动，相关筹备与执行分工由双方另行书面约定。")

add_heading_cn(doc, "（三）合作平台名称", level=2)
add_body(doc,
    "蔚蓝海洋—厦门大学闽投一号海洋科研合作基地（以下简称\"基地\"），"
    "并依据本协议第（二）项约定，在闽投一号同步挂牌为"
    "\"ONCE国际大科学计划·闽投一号创新科研基地\"。")

add_heading_cn(doc, "（四）合作方向", level=2)

add_heading_cn(doc, "1. 海洋环境智能监测", level=3)
add_body(doc,
    "依托乙方在海洋环境观测、海洋生态学、水声学、海洋传感与遥感、海洋大数据等领域的学科优势，"
    "结合甲方闽投一号现有平台及周边海域，联合开展水文、水质、生物、气象等多要素长期在线监测系统建设，"
    "形成可复用、可扩展的深远海智能监测体系。")

add_heading_cn(doc, "2. 水域净化与生态修复", level=3)
add_body(doc,
    "围绕养殖尾水处理、海域富营养化防控、海洋牧场底质改良、藻类与贝类生态净化等关键问题，"
    "联合开展技术研究与现场试验，形成可推广的水域净化与生态修复技术方案，"
    "提升闽投一号及周边海域的生态承载力与产业可持续性。")

add_heading_cn(doc, "3. 水下巡检与数据采集", level=3)
add_body(doc,
    "联合开展面向深远海养殖网箱的水下机器人（ROV/AUV）巡检、网衣结构健康监测、鱼群行为识别、"
    "水下成像与数据采集等科研合作，推动智能装备在闽投一号上的工程化应用，"
    "形成水下巡检与数据采集的标准化作业流程与数据规范。")

add_heading_cn(doc, "4. 海洋监测体系标准化与推广", level=3)
add_body(doc,
    "在闽投一号试点验证基础上，双方共同将所形成的监测规范、数据接口与作业流程进行标准化封装，"
    "推动其纳入福建省海洋经济相关政策与标准框架，并逐步向全省乃至全国推广，"
    "支撑相关海洋经济主体在合规接入、第三方核证、政府数据上报等环节的体系化对接。")

add_heading_cn(doc, "5. 数据资产化与高价值应用", level=3)
add_body(doc,
    "在长期监测数据持续沉淀的基础上，双方联合开展多方向的高价值应用研究，"
    "包括但不限于：蓝碳方法学研究与海洋碳汇核算，以监测数据为依据的绿色金融产品"
    "（绿色债券、海洋保险、ESG评估等）研究，面向海渔、生态环境等主管部门的长期数据服务方案研究。")

add_heading_cn(doc, "6. 联合申报国家级与省部级重点科研课题", level=3)
add_body(doc,
    "双方组织科研团队，围绕海洋强国战略以及福建省海洋经济发展相关政策导向，"
    "联合申报国家重点研发计划、国家自然科学基金、福建省科技重大专项等课题。"
    "联合课题的管理、经费分配及使用，按课题立项批复要求实施。")

add_heading_cn(doc, "7. 联合开展技术研究和产品开发", level=3)
add_body(doc,
    "发挥双方在深远海养殖装备、海洋观测仪器、水下机器人、海洋大数据与人工智能等领域的技术积累与产业资源，"
    "共同开展面向闽投一号及未来同类平台的技术研究与产品开发，"
    "并可根据需要拓展至其他海洋科技领域。")

add_heading_cn(doc, "（五）合作形式与内容", level=2)

add_heading_cn(doc, "1. 在校学生培养", level=3)
add_body(doc,
    "甲方为乙方师生提供闽投一号现场科研实习与实践基地，每年安排2周、4周、6周等时间不等的现场实习。"
    "实习内容包括但不限于：海洋环境监测实践、深远海养殖科学实践、水下装备测试与数据采集实践等。"
    "乙方围绕海洋经济与深远海产业关键技术，结合现场需求，编写教材教案、开展课程实践。")
add_body(doc,
    "甲方可为乙方在读研究生提供研究课题，并委派技术骨干作为校外导师；"
    "可为乙方在读研究生提供毕业设计指导。乙方优先推荐毕业生到甲方就业，"
    "甲方优先招聘双方联合培养的、符合招聘条件的研究生。")

add_heading_cn(doc, "2. 产业人才培训", level=3)
add_body(doc,
    "乙方依托海洋、生态、信息、管理等跨学科专业优势，面向甲方及福建省海洋经济相关企业从业人员，"
    "提供覆盖海洋观测、生态修复、智能装备、数据治理、蓝碳与绿色金融等多层次、跨学科的产业人才培训，"
    "助力甲方及海洋经济生态培养复合型、专业型人才。")

add_heading_cn(doc, "3. 学术研讨活动", level=3)
add_body(doc,
    "双方联合开展学术研讨、技术论坛、国际交流等学术活动，"
    "围绕深远海养殖、海洋牧场、蓝色碳汇、海洋数据要素等热点方向，"
    "联合承办学术会议、展示展览、竞赛比赛，共同编著相关书籍、发表论文等。")

add_heading_cn(doc, "（六）合作推进机制", level=2)
add_body(doc,
    "双方以紧密合作、加快推进为目标，分别指定对口联系部门，"
    "负责本协议项下合作事项的日常对接、统一协调与组织实施。"
    "甲方指定        作为对口联系部门，乙方指定国内合作办公室作为对口联系部门。"
    "本协议项下涉及的管理机构、议事规则及具体人员安排，由双方根据合作开展情况另行书面约定。")

# Section 2: Organizational structure (simplified - no committee / office)
add_heading_cn(doc, "二、组织机构", level=1)
add_body(doc,
    "1. 基地为非独立法人机构，由甲乙双方在本协议指导下共同管理和运作。")
add_body(doc,
    "2. 甲方依托闽投一号项目公司及相关业务部门开展工作；"
    "乙方依托海洋与地球学院、近海海洋环境科学国家重点实验室、环境与生态学院、"
    "生命科学学院、信息学院、国内合作办公室等开展工作。"
    "后续可根据需求经双方协商确定引入其他方参与共建。")
add_body(doc,
    "3. 基地的日常运行、管理机构设置、人员派遣及议事规则等事项，"
    "由双方根据合作推进情况另行书面约定，并可视情况以补充协议形式予以细化。")

# Section 3: Results and ownership
add_heading_cn(doc, "三、基地成果及归属", level=1)

add_heading_cn(doc, "（一）基地成果", level=2)
add_body(doc,
    "基地成果包括但不限于：设计资料、源代码、专利、软件著作权、学术论文、著作、成果原型、"
    "监测数据集与数据接口规范、技术标准与作业规范，以及基于联合研发成果所申报的各类课题。")

add_heading_cn(doc, "（二）基地成果归属与分享", level=2)
add_body(doc,
    "1. 甲乙双方签订本协议合作之前的知识产权归属不因本协议改变，依旧属于双方各自所有；")
add_body(doc,
    "2. 基地成果由甲乙双方共有；专利许可和转让需经甲乙双方书面同意，"
    "专利转让和许可收益分配双方协商确定；项目合作期间所取得的研究成果进行技术鉴定或申报奖励时，"
    "主出资的一方为第一完成单位；学术论文、著作等署名顺序根据实际贡献协商确定；")
add_body(doc,
    "3. 双方基于基地的研究成果对外发表论文，需经合作双方书面技术和保密审查同意，"
    "且在论文中须注明得到基地的支持；")
add_body(doc,
    "4. 基地在研究开发过程中产生的原始记录、监测数据、照片、录音、录像、文档等多种形式的真实记录，"
    "需妥善保管，未经双方书面同意，严禁向甲乙双方以外的第三方查阅、使用或泄露；")
add_body(doc,
    "5. 闽投一号现场所产生的监测数据，原则上由甲方持有并管理，乙方依研究目的依规使用；"
    "用于公开发表、对外披露或商业化应用前，应取得甲方书面同意。")

# Section 4: Organizational guarantee
add_heading_cn(doc, "四、组织保障", level=1)

add_heading_cn(doc, "（一）甲方责任", level=2)
add_body(doc,
    "1. 甲方将闽投一号深远海智能养殖平台、相关海域作业场地、配套设施及必要的现场作业资源"
    "开放用于基地的科研合作，作为甲方在本协议项下的主要投入；"
    "具体使用安排、作业窗口、安全保障与配套服务由双方在子协议中另行约定；")
add_body(doc,
    "2. 在甲方主营业务允许的范围内，向基地提供闽投一号平台的运营数据、生产数据"
    "及其他与合作方向相关的产业数据，支持联合科研工作的开展；")
add_body(doc, "3. 负责基地科研活动在闽投一号现场实施的统筹协调、安全管理与后勤保障；")
add_body(doc,
    "4. 配合乙方推动ONCE国际大科学计划在闽投一号的挂牌落地，"
    "提供挂牌仪式、品牌活动、对外宣传等所需的现场条件与资源支持；")
add_body(doc, "5. 将基地的项目成果及产业化进展定期向乙方反馈，并就工程化落地与产业化推广提供产业资源对接；")
add_body(doc, "6. 根据乙方对建设成果的意见，配合制订相应改进方案，完善成果；")
add_body(doc,
    "7. 明确甲方委派支持基地建设的技术骨干，并根据项目执行情况推荐联合培养的"
    "在读全日制和非全日制研究生，具体人数双方协商确定；")
add_body(doc,
    "8. 如双方就特定课题约定需追加专项经费或现金投入，由双方另行签订书面协议确定，"
    "并按相应协议管理使用。")

add_heading_cn(doc, "（二）乙方责任", level=2)
add_body(doc, "1. 负责本协议项下基地的科研规划、方案设计与学术质量管理；")
add_body(doc, "2. 负责对基地科研方案进行可行性论证，并提供专业技术、人才和资料支撑；")
add_body(doc,
    "3. 牵头对接ONCE国际大科学计划国际秘书处与相关国际科研体系，"
    "推动其在闽投一号挂牌落地及方法学本地化验证；")
add_body(doc, "4. 牵头组织联合申报国家级、省部级重点科研课题；")
add_body(doc, "5. 明确支持基地建设的技术专家团队，并为双方联合培养研究生提供研究课题；")
add_body(doc, "6. 负责联合研发成果的对外学术发布与共同推广。")

# Section 5: Confidentiality
add_heading_cn(doc, "五、保密", level=1)
add_body(doc,
    "1. 在基地的合作期间，甲乙双方对相关的国家秘密、技术秘密、商业秘密、闽投一号运营数据"
    "及科研监测数据均承担保密义务。未经对方书面许可，任何一方不得擅自披露对方的技术资料和商业信息，"
    "包括但不限于本协议的内容及与本协议有关的技术信息和经营信息；")
add_body(doc,
    "2. 任何一方由于违反保密条款而造成对方秘密信息泄露的，应承担由此引起的法律责任，"
    "并赔偿因此给对方造成的损失；")
add_body(doc, "3. 双方保密义务不因本协议的中止、撤销、解除、终止而终止。")

# Section 6: Term and others
add_heading_cn(doc, "六、协议期限及其他", level=1)
add_body(doc,
    "1. 本协议有效期三年，自本协议生效之日起算。本协议有效期届满前的六个月内，"
    "甲乙双方视合作情况，共同协商是否续签或终止本协议。")
add_body(doc, "2. 本协议的变更须由甲乙双方协商，并以书面形式确定。"
    "出现下列情形之一，致使本协议的履行成为不必要或不可能的，"
    "任何一方有权向另一方发出书面通知提前终止本协议：")
add_body(doc, "    （1）因发生不可抗力或重大技术风险；")
add_body(doc, "    （2）另一方实质性违反本协议的约定，并经书面催告后30天内未对该违约行为进行补救；")
add_body(doc, "    （3）本协议一旦终止，各方均应立即返还自对方接收的全部保密信息及其副本。")
add_body(doc, "3. 甲乙双方因履行本协议过程中出现的争议，按照下述方式处理：")
add_body(doc, "    （1）本协议的未尽事宜，双方本着坦诚合作的原则友好协商解决；")
add_body(doc,
    "    （2）在协议履行过程中发生任何争议，双方本着互谅互让的原则协商解决，"
    "如若协商未果，双方同意任何一方均可向甲方住所地人民法院提起诉讼；")
add_body(doc, "    （3）甲乙双方应严格履行本协议约定，如有违反，违约方必须承担相应的经济责任和法律责任。")

# Section 7: Supplementary
add_heading_cn(doc, "七、附则", level=1)
add_body(doc,
    "1. 本协议一式肆份，由甲乙双方法定代表人或授权代表正式签署盖章后生效，"
    "双方各持贰份，具有同等的法律效力。")
add_body(doc,
    "2. 在本协议有效期内，甲方指定          为甲方联系人，联系方式          ；"
    "乙方指定          为乙方联系人，联系方式          。"
    "一方变更联系人的，应当及时以书面形式通知另一方。")
add_body(doc,
    "3. 本协议构成甲乙双方关于本协议所述主题的完整协议，"
    "并取代所有之前关于本协议所述主题的任何协议及讨论。"
    "未经双方事先书面同意，任一方不得对本协议进行任何变更，"
    "也不得擅自转让本协议约定的任何权利或义务。")
add_body(doc,
    "4. 本协议的未尽事宜，由甲乙双方友好协商解决并签订补充协议，"
    "补充协议与本协议具有同等的法律效力。")
add_body(doc, "5. 以下无正文。", indent=False)

doc.add_page_break()

# Signing page
add_heading_cn(doc, "签 署 页", level=1)
p = doc.paragraphs[-1]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_blank(doc)
add_body(doc, "甲方：蔚蓝海洋集团有限公司（盖章）", indent=False)
add_body(doc, "法定代表人/授权代表：", indent=False)
add_body(doc, "联系人：", indent=False)
add_body(doc, "地址：", indent=False)
add_body(doc, "联系电话：               传真：               ", indent=False)
add_body(doc, "日期：二〇二六年    月    日", indent=False)

add_blank(doc)
add_blank(doc)

add_body(doc, "乙方：厦门大学（盖章）", indent=False)
add_body(doc, "法定代表人/授权代表：", indent=False)
add_body(doc, "联系人：", indent=False)
add_body(doc, "地址：福建省厦门市思明区思明南路422号", indent=False)
add_body(doc, "联系电话：               传真：               ", indent=False)
add_body(doc, "日期：二〇二六年    月    日", indent=False)

import os
out_path = os.path.join(os.path.dirname(__file__),
                       "蔚蓝海洋集团-厦门大学-闽投一号海洋科研合作框架协议（草案）.docx")
doc.save(out_path)
print("Generated:", out_path)
